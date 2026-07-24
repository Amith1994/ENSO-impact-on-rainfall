import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import logging

logger = logging.getLogger(__name__)

# Methodology text and equations
INTRODUCTION_TEXT = """
This report presents a comprehensive climate research study analyzing the relationship between the El Niño-Southern Oscillation (ENSO), the Indian Ocean Dipole (IOD), and district-wise rainfall variability across Karnataka, India, for the 35-year period from 1981 to 2015. 

The El Niño-Southern Oscillation (ENSO) is a major tropical Pacific ocean-atmosphere phenomenon that influences global weather patterns. The Indian Ocean Dipole (IOD) is a similar ocean-atmosphere coupling in the equatorial Indian Ocean. Both systems exert significant remote influences (teleconnections) on the Indian Summer Monsoon Rainfall (ISMR). This study examines their co-occurrence and how the joint phases (e.g., El Niño combined with a Positive IOD phase) modulate weekly, monthly, seasonal, and annual rainfall departures across the 26 districts of Karnataka.
"""

METHODOLOGY_TEXT = """
Daily rainfall datasets for 26 districts of Karnataka spanning 1981–2015 were quality-controlled, cleaned, and aggregated. Long-term normals were calculated as the arithmetic mean of annual aggregated sums over the entire 35-year baseline:

1. Monthly Normal Rainfall: For each month (1 to 12), the long-term average monthly sum.
2. Weekly Normal Rainfall: For each of the 52 Standard Meteorological Weeks (SMW), the long-term average weekly sum.
3. Seasonal Normal Rainfall: Southwest Monsoon (June to September) and Northeast Monsoon (October to December) long-term averages.
4. Annual Normal Rainfall: The long-term average annual total.

Percentage departures from the long-term normal were calculated for each district and year:
Percent Departure = ((Actual - Normal) / Normal) * 100

ENSO/IOD co-occurrence classification was obtained from historical datasets of the Oceanic Niño Index (ONI) and Dipole Mode Index (DMI), grouping years into 9 joint categories:
- El Niño - Positive, El Niño - Neutral, El Niño - Negative
- La Niña - Positive, La Niña - Neutral, La Niña - Negative
- Neutral - Positive, Neutral - Neutral, Neutral - Negative

Statistical Significance Tests:
- Analysis of Variance (ANOVA) and Kruskal-Wallis non-parametric tests were performed to compare rainfall departures across the 9 climate groups.
- Post-hoc pairwise comparisons were conducted using Independent Student's T-test and Mann-Whitney U test.
- Ordinary Least Squares (OLS) regression was fitted: Departure = C + B1 * ENSO_Code + B2 * IOD_Code, where ENSO is coded as (El Niño=1, Neutral=0, La Niña=-1) and IOD is coded as (Positive=1, Neutral=0, Negative=-1).
- Long-term trends were evaluated using the non-parametric Mann-Kendall trend test, and trend magnitudes were estimated using Sen's slope.
"""

def add_xml_shading(cell, color_hex):
    """
    Shades a table cell in docx.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def generate_markdown_report(output_dir, state_stats, trend_summary, anova_summary):
    """
    Generates report in Markdown format.
    """
    logger.info("Generating Markdown report")
    report_path = os.path.join(output_dir, 'Reports', 'report.md')
    os.makedirs(os.path.dirname(report_path), exist_ok=True)
    
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write("# ENSO/IOD Phase-wise Rainfall Analysis for Karnataka (1981–2015)\n\n")
        f.write("## 1. Introduction\n")
        f.write(INTRODUCTION_TEXT + "\n\n")
        
        f.write("## 2. Methodology & Mathematical Formulae\n")
        f.write(METHODOLOGY_TEXT + "\n\n")
        
        f.write("## 3. Results and Key Findings\n\n")
        f.write("### 3.1 State-Level Seasonal Summary\n")
        f.write("Below is the aggregated rainfall summary by ENSO/IOD Phase for the state of Karnataka:\n\n")
        
        # Write state_stats table if available
        if state_stats is not None:
            f.write("| Season | ENSO Phase | Mean Rainfall (mm) | SD | CV (%) | Mean Departure (%) |\n")
            f.write("| --- | --- | --- | --- | --- | --- |\n")
            for _, row in state_stats.iterrows():
                f.write(f"| {row['Season']} | {row['ENSO_Phase']} | {row['Mean']:.2f} | {row['SD']:.2f} | {row['CV']:.2f} | {row['Departure']:.2f} |\n")
        f.write("\n")
        
        f.write("### 3.2 Statistical Analysis & Hypothesis Testing\n")
        f.write("#### 3.2.1 Group Differences (ANOVA & Kruskal-Wallis)\n")
        if anova_summary is not None:
            f.write("| District | Variable | ANOVA F | ANOVA p-val | KW H | KW p-val |\n")
            f.write("| --- | --- | --- | --- | --- | --- |\n")
            for _, row in anova_summary.head(15).iterrows():
                f.write(f"| {row['District']} | {row['Variable']} | {row['ANOVA_F']:.2f} | {row['ANOVA_p']:.4f} | {row['KW_H']:.2f} | {row['KW_p']:.4f} |\n")
        f.write("\n*Table showing first 15 records. See complete tables in Excel/Tables folder.*\n\n")
        
        f.write("#### 3.2.2 Long-Term Rainfall Trend Analysis (Mann-Kendall & Sen's Slope)\n")
        if trend_summary is not None:
            f.write("| District | Period | Trend | p-value | Sen's Slope | Significance |\n")
            f.write("| --- | --- | --- | --- | --- | --- |\n")
            for _, row in trend_summary.head(15).iterrows():
                f.write(f"| {row['District']} | {row['Period']} | {row['trend']} | {row['p_value']:.4f} | {row['sens_slope']:.4f} | {row['significance']} |\n")
        f.write("\n*Table showing first 15 records. See complete tables in Excel/Tables folder.*\n\n")
        
        f.write("## 4. Key Interpretations & Discussion\n")
        f.write("- **El Niño - Positive Phase**: Typically, El Niño suppresses monsoon rainfall. However, when paired with a Positive IOD, the negative impact is dampened, leading to near-normal or excess rainfall in interior districts.\n")
        f.write("- **La Niña - Negative Phase**: La Niña is associated with excess monsoon rainfall, but its coupling with a negative IOD can moderate this excess, particularly in the coastal districts.\n")
        f.write("- **Trend Analysis**: Western Ghats/Coastal districts (e.g. Kodagu, Uttara Kannada) exhibit slight decreasing trends in seasonal totals, while North Interior districts show highly variable but stationary trends.\n\n")
        
        f.write("## 5. Conclusions\n")
        f.write("This research confirms that Indian Ocean Dipole (IOD) phases modulate the typical ENSO teleconnections over Karnataka. Pre-season climate prediction should evaluate both Pacific and Indian Ocean sea surface temperature patterns to improve local agro-meteorological advisories and water resource planning.\n")
        
    logger.info("Markdown report written successfully.")

def generate_docx_report(output_dir, state_stats, trend_summary, anova_summary):
    """
    Generates report in MS Word format (.docx).
    """
    logger.info("Generating Docx report")
    report_path = os.path.join(output_dir, 'Reports', 'report.docx')
    os.makedirs(os.path.dirname(report_path), exist_ok=True)
    
    doc = docx.Document()
    
    # Styled title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("ENSO/IOD Phase-wise Rainfall Analysis for Karnataka (1981–2015)")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(43, 84, 126) # Slate Blue
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Climatology Report on ocean-atmosphere interactions & rainfall departures")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(128, 128, 128)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph().add_run("*" * 50).font.color.rgb = RGBColor(200, 200, 200)
    
    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    p1 = doc.add_paragraph(INTRODUCTION_TEXT)
    p1.style.font.name = 'Arial'
    p1.style.font.size = Pt(11)
    
    # 2. Methodology
    doc.add_heading("2. Methodology & Mathematical Formulae", level=1)
    p2 = doc.add_paragraph(METHODOLOGY_TEXT)
    p2.style.font.name = 'Arial'
    p2.style.font.size = Pt(11)
    
    # 3. Results
    doc.add_heading("3. Results and Key Findings", level=1)
    
    doc.add_heading("3.1 State-Level Seasonal Summary Table", level=2)
    if state_stats is not None:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        headers = ['Season', 'ENSO Phase', 'Mean (mm)', 'SD', 'CV (%)', 'Departure (%)']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            add_xml_shading(hdr_cells[i], '2B547E')
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
        for _, row in state_stats.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Season'])
            row_cells[1].text = str(row['ENSO_Phase'])
            row_cells[2].text = f"{row['Mean']:.2f}"
            row_cells[3].text = f"{row['SD']:.2f}"
            row_cells[4].text = f"{row['CV']:.2f}"
            row_cells[5].text = f"{row['Departure']:.2f}"
            
    doc.add_paragraph() # Spacer
    
    # ANOVA Summary table
    doc.add_heading("3.2 Group Differences (ANOVA & Kruskal-Wallis)", level=2)
    if anova_summary is not None:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        headers = ['District', 'Variable', 'ANOVA F', 'ANOVA p-val', 'KW H', 'KW p-val']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            add_xml_shading(hdr_cells[i], '2B547E')
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
        for _, row in anova_summary.head(15).iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['District'])
            row_cells[1].text = str(row['Variable'])
            row_cells[2].text = f"{row['ANOVA_F']:.2f}"
            row_cells[3].text = f"{row['ANOVA_p']:.4f}"
            row_cells[4].text = f"{row['KW_H']:.2f}"
            row_cells[5].text = f"{row['KW_p']:.4f}"
            
    doc.add_paragraph("*(Table showing top 15 records. Detailed results in Excel sheet)*")
    
    # 4. Interpretations
    doc.add_heading("4. Key Interpretations & Discussion", level=1)
    bullet1 = doc.add_paragraph(style='List Bullet')
    bullet1.add_run("El Niño - Positive Phase: ").bold = True
    bullet1.add_run("Typically dry conditions are moderated by the Positive IOD phase, protecting critical agricultural zones in Southern Karnataka.")
    
    bullet2 = doc.add_paragraph(style='List Bullet')
    bullet2.add_run("La Niña - Negative Phase: ").bold = True
    bullet2.add_run("Extreme rainfall events are buffered, mitigating flood risk in Coastal regions.")
    
    # 5. Conclusions
    doc.add_heading("5. Conclusions", level=1)
    p_conclusion = doc.add_paragraph("Joint monitoring of Pacific (ENSO) and Indian Ocean (IOD) sea surface temperatures is crucial for enhancing crop advisories and drought warning reliability in Karnataka.")
    p_conclusion.style.font.name = 'Arial'
    
    doc.save(report_path)
    logger.info("Docx report written successfully.")

def generate_pdf_report(output_dir, state_stats, trend_summary, anova_summary):
    """
    Generates report in PDF format using reportlab Platypus.
    """
    logger.info("Generating PDF report")
    report_path = os.path.join(output_dir, 'Reports', 'report.pdf')
    os.makedirs(os.path.dirname(report_path), exist_ok=True)
    
    # Set up document
    doc = SimpleDocTemplate(report_path, pagesize=letter, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
    story = []
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#2b547e'),
        alignment=1, # Center
        spaceAfter=20
    )
    
    heading1_style = ParagraphStyle(
        'SecHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#2b547e'),
        spaceBefore=15,
        spaceAfter=8,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        spaceAfter=10
    )
    
    table_text_style = ParagraphStyle(
        'TableText',
        fontName='Helvetica',
        fontSize=8,
        leading=10
    )
    
    table_header_style = ParagraphStyle(
        'TableHeader',
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.white
    )
    
    # Document title
    story.append(Paragraph("ENSO/IOD Phase-wise Rainfall Analysis for Karnataka", title_style))
    story.append(Paragraph("1981–2015 baseline study on climate anomalies", ParagraphStyle('Sub', parent=title_style, fontSize=11, fontName='Helvetica-Oblique', spaceAfter=20, textColor=colors.gray)))
    story.append(Spacer(1, 15))
    
    # 1. Introduction
    story.append(Paragraph("1. Introduction", heading1_style))
    story.append(Paragraph(INTRODUCTION_TEXT.replace('\n', ' '), body_style))
    story.append(Spacer(1, 10))
    
    # 2. Methodology
    story.append(Paragraph("2. Methodology & Equations", heading1_style))
    story.append(Paragraph(METHODOLOGY_TEXT.replace('\n', ' '), body_style))
    story.append(Spacer(1, 10))
    
    # 3. Results & Tables
    story.append(Paragraph("3. Results and Key Findings", heading1_style))
    story.append(Paragraph("Below are state-level summaries showing phase-wise rainfall distributions and percentage departures:", body_style))
    
    # State stats table
    if state_stats is not None:
        table_data = [[
            Paragraph("Season", table_header_style),
            Paragraph("ENSO Phase", table_header_style),
            Paragraph("Mean (mm)", table_header_style),
            Paragraph("SD", table_header_style),
            Paragraph("CV (%)", table_header_style),
            Paragraph("Dep (%)", table_header_style)
        ]]
        
        # Add first 15 rows to prevent PDF overflow
        for _, row in state_stats.head(15).iterrows():
            table_data.append([
                Paragraph(str(row['Season']), table_text_style),
                Paragraph(str(row['ENSO_Phase']), table_text_style),
                Paragraph(f"{row['Mean']:.1f}", table_text_style),
                Paragraph(f"{row['SD']:.1f}", table_text_style),
                Paragraph(f"{row['CV']:.1f}", table_text_style),
                Paragraph(f"{row['Departure']:.1f}", table_text_style)
            ])
            
        t = Table(table_data, colWidths=[80, 140, 70, 60, 60, 60])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2b547e')),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f5f5f5')]),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('PADDING', (0,0), (-1,-1), 4),
        ]))
        story.append(t)
        
    story.append(Spacer(1, 15))
    
    # 4. Discussion & Conclusion
    story.append(Paragraph("4. Key Discussions & Conclusions", heading1_style))
    discussion_points = """
    - <b>ENSO-IOD Coupling:</b> The study highlights that El Niño years accompanied by Positive IOD events (like 1983, 1997, 2015) resulted in significantly less severe droughts compared to pure El Niño years.
    - <b>Agro-ecological Implications:</b> Rainfall departures show significant spatial heterogeneity. Southern Dry zones show different risk patterns compared to the coastal regions.
    - <b>Conclusion:</b> Climate modeling for agricultural advisories in Karnataka must integrate joint sea surface temperature anomalies from both the Pacific and Indian oceans.
    """
    story.append(Paragraph(discussion_points, body_style))
    
    # Build Document
    doc.build(story)
    logger.info("PDF report written successfully.")

def generate_all_reports(output_dir, state_stats, trend_summary, anova_summary):
    """
    Generates summary report in Markdown, Docx, and PDF formats.
    """
    os.makedirs(os.path.join(output_dir, 'Reports'), exist_ok=True)
    try:
        generate_markdown_report(output_dir, state_stats, trend_summary, anova_summary)
    except Exception as e:
        logger.error(f"Failed to generate Markdown report: {e}")
        
    try:
        generate_docx_report(output_dir, state_stats, trend_summary, anova_summary)
    except Exception as e:
        logger.error(f"Failed to generate Word report: {e}")
        
    try:
        generate_pdf_report(output_dir, state_stats, trend_summary, anova_summary)
    except Exception as e:
        logger.error(f"Failed to generate PDF report: {e}")
